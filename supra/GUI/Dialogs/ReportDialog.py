import os
import numpy as np
import matplotlib.pyplot as plt
import pickle
from datetime import datetime, timedelta
now_time = datetime.now()
from functools import partial

from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
import pyqtgraph as pg

from supra.GUI.Tools.Theme import theme
from supra.GUI.Tools.GUITools import *
from supra.GUI.Tools.ReportHelper import *
from supra.Utils.AngleConv import chauvenet
from supra.Supracenter.anglescanrev import anglescanrev
from supra.Supracenter.cyscanVectors import cyscan as cyscanV
from supra.Supracenter.cyscan5 import cyscan

from supra.GUI.Tabs.SupracenterSearch import supSearch, resultsPrint

from supra.Files.SaveObjs import Prefs

from supra.Utils.Classes import Position

from supra.Stations.Filters import *


try:
    from docx import Document
    docx_import = True
except:
    docx_import = False

SPREAD = 3

def makePropLine(ref_pos, D, alpha=255):

    lats = []
    lons = []
    for line in D:
        temp = Position(0, 0, 0)
        temp.x = line[0]
        temp.y = line[1]
        temp.z = line[2]
        temp.pos_geo(ref_pos)
        if not np.isnan(temp.lat) and not np.isnan(temp.lon):
            lats.append(temp.lat)
            lons.append(temp.lon)

    lats.sort()
    lons.sort()

    return lats, lons

def propegateBack(bam, stn, azimuth, offset=0, frag_height=30000):

    ref_pos = Position(bam.setup.lat_centre, bam.setup.lon_centre, 0)

    S = stn.metadata.position
    
    S.pos_loc(ref_pos)

    sounding, perturbations = bam.atmos.getSounding(lat=[S.lat, S.lat], lon=[S.lon, S.lon], heights=[S.elev, frag_height])

    D = []
    # Use 25 angles between 90 and 180 deg

    ### To do this more right, calculate D with bad winds, and then use D to find new winds and then recalc D

    for zenith in np.linspace(1, 89, 25):
        # D = anglescanrev(S.xyz, self.azimuth + offset, zenith, sounding, wind=True)
        # D = anglescanrev(S.xyz, (self.azimuth + offset + 180)%360, zenith, sounding, wind=True)

        D.append(anglescanrev(S.xyz, azimuth + offset, zenith, sounding, wind=True))
        D.append(anglescanrev(S.xyz, (azimuth + offset + 180)%360, zenith, sounding, wind=True))
    # pt, err = finalanglecheck(self.bam, self.bam.setup.trajectory, self.stn.metadata.position, self.azimuth)

    start_pt = makePropLine(ref_pos, np.array(D))


    return start_pt


class ReportWindow(QWidget):

    def __init__(self, bam, prefs):

        QWidget.__init__(self)
        
        self.bam = bam
        self.prefs = prefs

        plt.style.use('default')
        fig = plt.figure()

        self.buildGUI()

    def buildGUI(self):
        self.setWindowTitle('Generate Report')
        app_icon = QtGui.QIcon()
        app_icon.addFile(os.path.join('supra', 'GUI', 'Images', 'preferences.png'), QtCore.QSize(16, 16))
        self.setWindowIcon(app_icon)

        p = self.palette()
        p.setColor(self.backgroundRole(), Qt.black)
        self.setPalette(p)

        theme(self)

        layout = QGridLayout()
        self.setLayout(layout)

        self.general_info = QCheckBox('General Information')
        layout.addWidget(self.general_info, 0, 1)

        self.def_info = QCheckBox('Definitions/Methods')
        layout.addWidget(self.def_info, 1, 1)

        self.traj_info = QCheckBox('Trajectory')
        layout.addWidget(self.traj_info, 2, 1)

        self.frag_info = QCheckBox('Fragmentation(s)')
        layout.addWidget(self.frag_info, 3, 1)

        self.stat_info = QCheckBox('Stations')
        layout.addWidget(self.stat_info, 4, 1)

        self.atmos_info = QCheckBox('Atmosphere')
        layout.addWidget(self.atmos_info, 5, 1)

        self.supra_info = QCheckBox('Supracenter Solution')
        layout.addWidget(self.supra_info, 6, 1)

        save_button = QPushButton('Generate Report')
        layout.addWidget(save_button, 8, 2)
        save_button.clicked.connect(self.genRep)
     
    def genRep(self):

        if docx_import:
            print("Printing Report...")

            file_name = fileSearch(['docx (*.docx)'], None)

            doc = Document()

            try:
                doc.add_heading('METEOR REPORT: {:}'.format(self.bam.setup.fireball_name), 0)
            except AttributeError:
                doc.add_heading('METEOR REPORT', 0)   
                doc.add_paragraph('No fireball name detected -> Possible error in document, proceed with caution!')

            if self.general_info.isChecked():
                print("... General Information")
                doc.add_heading('General Information', 1)
                self.writeEvtPreamble(doc)

            if self.def_info.isChecked():
                print("... Definitions/Methods")
                doc.add_heading('Definitions/Methods', 1)
                self.writeDefs(doc)

            if self.traj_info.isChecked():
                try:
                    print("... Trajectory")
                    doc.add_heading('Trajectory')
                    self.writeTraj(doc)
                except:
                    pass

            if self.frag_info.isChecked():
                print("... Fragmentation(s)")
                doc.add_heading('Fragmentation(s)')
                self.writeFrags(doc)

            if self.stat_info.isChecked():
                print("... Stations")
                doc.add_heading('Stations', 1)
                self.writeStats(doc)

            if self.atmos_info.isChecked():
                print("... Atmosphere")
                doc.add_heading('Atmosphere', 1)
                self.writeAtm(doc)

            if self.supra_info.isChecked():
                print("... Supracenter Solution")
                doc.add_heading('Supracenter Solution', 1)
                self.writeSupSol(doc)


            doc.save(file_name)

        else:
            print('python-docx not detected: I`m not generating anything now')

        self.close()

    def writeTraj(self, doc):

        traj = self.bam.setup.trajectory
        src = self.bam.setup.traj_metadata[0]

        doc.add_paragraph('Title: {:}'.format(src.title))
        doc.add_paragraph('Notes: {:}'.format(src.notes))
        doc.add_paragraph('Azimuth: {:}°,  Zenith: {:}°'.format(traj.azimuth.deg, traj.zenith.deg))
        doc.add_paragraph('Time:    {:}s,  Average Velocity: {:} km/s'.format(traj.t, traj.v/1000))
        doc.add_paragraph('Initial Position:')
        doc.add_paragraph('{:}, {:}, {:.2f} m'.format(latitudify(traj.pos_i.lat), \
                                                          longitudify(traj.pos_i.lon), \
                                                          traj.pos_i.elev))
        doc.add_paragraph('Final Position:')
        doc.add_paragraph('{:}, {:}, {:.2f} m'.format(latitudify(traj.pos_f.lat), \
                                                             longitudify(traj.pos_f.lon), \
                                                             traj.pos_f.elev))

    def writeFrags(self, doc):
        
        for i, frag in enumerate(self.bam.setup.fragmentation_point):

            src = self.bam.setup.frag_metadata[i]

            doc.add_paragraph('Title: {:}'.format(src.title))
            doc.add_paragraph('Notes: {:}'.format(src.notes))
            doc.add_paragraph('Fragmentation {:}: {:}, {:}, {:.2f} km, {:}'.format(i+1, latitudify(frag.position.lat), \
                                                             longitudify(frag.position.lon), \
                                                             frag.position.elev/1000, self.bam.setup.fireball_datetime + timedelta(seconds=frag.time)))

    def writeEvtPreamble(self, doc):

        lat = latitudify(self.bam.setup.lat_centre)
        lon = longitudify(self.bam.setup.lon_centre)
        file_size = byteify(os.stat(self.bam.file_name).st_size)

        doc.add_paragraph('This document was produced by the Bolide Acoustic Modelling (BAM) program created by the Western Meteor Physics Group at the University of Western Ontario.')
        doc.add_paragraph('Document Creation: {:} [local time]'.format(now_time))

        doc.add_paragraph('Approximate Location: {:} {:}'.format(lat, lon))
        doc.add_paragraph('Time of first sighting: {:} UTC'.format(self.bam.setup.fireball_datetime))
        doc.add_paragraph('BAM File: {:} ({:})'.format(self.bam.file_name, file_size))


    def writeDefs(self, doc):

        defdir = os.path.join("supra", "Misc", "definitions.txt")

        with open(defdir, "r+") as f:

            lines = [line.rstrip().split(',') for line in f]

        for d in lines:
            doc.add_heading('{:}'.format(d[0]), 2)
            doc.add_paragraph('{:}'.format(d[1]))


    def writeStats(self, doc):

        doc.add_heading('Map', 2)
        self.makeStationMap(doc)


        for stat in self.bam.stn_list:
            doc.add_heading('-----{:}-{:}-----'.format(stat.metadata.network, stat.metadata.code), 2)
            doc.add_heading('Metadata', 3)
            doc.add_paragraph('Name: {:}'.format(stat.metadata.name))
            try:
                doc.add_paragraph('Bandpass: {:}-{:} Hz'.format(stat.metadata.low_bandstop, stat.metadata.high_bandstop))
            except AttributeError:
                print("Station has no bandpass")
            doc.add_paragraph('Location: {:}, {:}, {:.2f} m'.format(latitudify(stat.metadata.position.lat), \
                                                             longitudify(stat.metadata.position.lon), \
                                                             stat.metadata.position.elev))
            doc.add_paragraph('Obtained from: {:}'.format(stat.metadata.source))

            if len(stat.polarization.azimuth) > 0:
                doc.add_heading('Polarization', 3)
                poldata = []
                for pol in range(len(stat.polarization.azimuth)):

                    doc.add_paragraph('{:.2f} ± {:.2f}°'.format(stat.polarization.azimuth[pol], stat.polarization.azimuth_error[pol]))

                    # Using pols in same order as frags -> might want to fix this

                    az = stat.polarization.azimuth[pol]
                    frag_h = self.bam.setup.fragmentation_point[0].position.elev
                    az_err = stat.polarization.azimuth_error[pol]

                    nomlat, nomlon = propegateBack(self.bam, stat, az, frag_height=frag_h)
                    maxlat, maxlon = propegateBack(self.bam, stat, az, offset=az_err, frag_height=frag_h)
                    minlat, minlon = propegateBack(self.bam, stat, az, offset=-az_err, frag_height=frag_h)

                poldata.append([nomlat, nomlon, maxlat, maxlon, minlat, minlon])



                for stat in self.bam.stn_list:
                    plt.scatter(stat.metadata.position.lon, stat.metadata.position.lat, c='r', marker='^')
                    plt.text(stat.metadata.position.lon, stat.metadata.position.lat, '{:}-{:}'.format(stat.metadata.network, stat.metadata.code))

                plt.plot([self.bam.setup.trajectory.pos_i.lon, self.bam.setup.trajectory.pos_f.lon],\
                                                      [self.bam.setup.trajectory.pos_i.lat, self.bam.setup.trajectory.pos_f.lat])
                plt.scatter([self.bam.setup.trajectory.pos_f.lon], \
                                                             [self.bam.setup.trajectory.pos_f.lat], \
                                                                marker='x')

                for polline in poldata:
                    plt.plot(polline[1], polline[0])
                    plt.plot(polline[3], polline[2])
                    plt.plot(polline[5], polline[4])

                plt.xlabel('Longitude')
                plt.ylabel('Latitude')
                pic_file = os.path.join(self.prefs.workdir, self.bam.setup.fireball_name, 'Station_pol_map.png')
                
                plt.savefig(pic_file)
                doc.add_picture(pic_file)
                os.remove(pic_file)


            doc.add_heading('Ballistic Arrivals', 3)
            b_times = []
            b_time_err = []
            try:
                ball_time = stat.times.ballistic[0][0][0]
                ball_perts, _ = chauvenet(stat.times.ballistic[0][1][0])
                ball_max = np.abs(np.nanmax(ball_perts) - ball_time)
                ball_min = np.abs(np.nanmin(ball_perts) - ball_time)

                if ball_max != ball_min:
                    doc.add_paragraph('{:.2f} +{:.2f} / -{:.2f} s'.format(ball_time, ball_max, ball_min))
                else:
                    doc.add_paragraph('{:.2f} ± {:.2f} s'.format(ball_time, ball_max))
                b_times.append(ball_time)
                b_time_err.append([[ball_min], [ball_max]])

            except IndexError:
                doc.add_paragraph('No ballistic arrival')
            except ValueError:

                # TODO add proper handling here
                doc.add_paragraph('No ballistic arrival')

            self.makeWaveform(stat, b_times, b_time_err, doc, typ='ball')
            
            doc.add_heading('Fragmentation Arrivals', 3)
            f_times = []
            f_time_err = []
            for i, frag in enumerate(self.bam.setup.fragmentation_point):
                try:
                    f_time = stat.times.fragmentation[i][0][0]
                    f_perts, _ = obtainPerts(stat.times.fragmentation, i)

                    if np.isnan(f_perts.all()):
                        raise AllNanError

                    f_max = np.abs(np.nanmax(f_perts) - f_time)
                    f_min = np.abs(np.nanmin(f_perts) - f_time)
                    if f_max != f_min:
                        doc.add_paragraph('Fragmentation {:} (F{:}): {:.2f} +{:.2f} / -{:.2f} s'.format(i+1, i+1, f_time, f_max, f_min))
                    else:
                        doc.add_paragraph('Fragmentation {:} (F{:}): {:.2f} ± {:.2f} s'.format(i+1, i+1, f_time, f_min))
                    f_times.append(f_time)
                    f_time_err.append([[f_min], [f_max]])

                except IndexError:
                    doc.add_paragraph('No Fragmentation {:} (F{:}) Arrival'.format(i+1, i+1))
                    f_times.append(np.nan)
                    f_time_err.append([[np.nan], [np.nan]]) 

                except AllNanError:
                    doc.add_paragraph('Fragmentation {:} (F{:}): {:.2f} s'.format(i+1, i+1, f_time))
                    f_times.append(f_time)
                    f_time_err.append([[np.nan], [np.nan]]) 


            self.makeWaveform(stat, f_times, f_time_err, doc, typ='frag')
            self.makeStaff(stat, f_times, f_time_err, doc)

            doc.add_heading('Precursor Arrivals', 3)
            p_times = []
            p_time_err = []
            for i, frag in enumerate(self.bam.setup.fragmentation_point):
                v_time = (frag.position.elev - stat.metadata.position.elev)/310
                h_time = frag.position.ground_distance(stat.metadata.position)/2000
                h_time_p = frag.position.ground_distance(stat.metadata.position)/4000
                h_time_m = frag.position.ground_distance(stat.metadata.position)/1000

                p_time = h_time + v_time
                p_time_p = np.abs(h_time_p + v_time - p_time)
                p_time_m = np.abs(h_time_m + v_time - p_time)
   
                if p_time_p != p_time_m:
                    doc.add_paragraph('Precursor {:} (P{:}): {:.2f} +{:.2f} / -{:.2f} s'.format(i+1, i+1, p_time, p_time_m, p_time_p))
                else:
                    doc.add_paragraph('Precursor {:} (P{:}): {:.2f} ± {:.2f} s'.format(i+1, i+1, p_time, p_time_m))

                p_times.append(p_time)
                p_time_err.append([[p_time_p], [p_time_m]])

            self.makeWaveform(stat, p_times, p_time_err, doc, typ='prec')

            times = f_times + b_times + p_times
            time_err = f_time_err + b_time_err + p_time_err

            divs = [len(f_times), len(b_times), len(p_times)]

            self.makeWaveform(stat, times, time_err, doc, typ='preview', divs=divs)

    def makeStationMap(self, doc):


        for stat in self.bam.stn_list:
            plt.scatter(stat.metadata.position.lon, stat.metadata.position.lat, c='r', marker='^')
            plt.text(stat.metadata.position.lon, stat.metadata.position.lat, '{:}-{:}'.format(stat.metadata.network, stat.metadata.code))

        plt.xlabel('Longitude')
        plt.ylabel('Latitude')
        pic_file = os.path.join(self.prefs.workdir, self.bam.setup.fireball_name, 'Station_map.png')
        
        plt.savefig(pic_file)
        doc.add_picture(pic_file)
        os.remove(pic_file)
        plt.clf()


    def makeWaveform(self, stn, times, time_err, doc, typ='frag', divs=[]):

        if len(times) == 0:
            return None

        st = stn.stream

        chn_list = []            

        for i in range(len(st)):
            chn = st[i].stats.channel[:2]
            
            chn_list.append((chn + '*'))

        chn_list = list(set(chn_list))

        doc.add_paragraph('Available Channels:')
        doc.add_paragraph(', '.join(chn_list))

        if 'BD*' in chn_list:
            chn_selected = 'BDF'
        elif 'HH*' in chn_list:
            chn_selected = 'HHZ'
        elif 'BH*' in chn_list:
            chn_selected = 'BHZ'

        st = st.select(channel=chn_selected)[0]
        st.detrend()
        waveform_data = st.data
        time_data = np.arange(0, st.stats.npts/st.stats.sampling_rate, st.stats.delta)

        waveform_data = waveform_data[:len(time_data)]
        start_datetime = st.stats.starttime.datetime
        end_datetime = st.stats.endtime.datetime
        stn.offset = (start_datetime - self.bam.setup.fireball_datetime).total_seconds()
        time_data = time_data[:len(waveform_data)] + stn.offset

        # Init the butterworth bandpass filter
        butter_b, butter_a = butterworthBandpassFilter(2, 8, \
            1.0/st.stats.delta, order=6)

        # Filter the data
        waveform_data = scipy.signal.filtfilt(butter_b, butter_a, np.copy(waveform_data))


        plt.plot(time_data, waveform_data)

        for tt in range(len(times)):

            x1,x2,y1,y2 = plt.axis()
            y_range = y2 - y1
            y_pos = 0.1*y_range*tt
            if typ == 'frag':
                title = 'Fragmentation Arrivals'
                plt.annotate('F{:}'.format(tt + 1), (times[tt], y_pos), zorder=4)
                plt.errorbar([times[tt]], [y_pos], xerr=time_err[tt], fmt='o', zorder=3)
            elif typ == 'ball':
                title = 'Ballistic Arrivals'
                plt.annotate('B ', (times[tt], y_pos), zorder=4)
                plt.errorbar([times[tt]], [y_pos], xerr=time_err[tt], fmt='o', zorder=3)
            elif typ == 'prec':
                title = 'Precursor Arrivals'
                plt.annotate('P{:}'.format(tt + 1), (times[tt], y_pos), zorder=4)
                plt.errorbar([times[tt]], [y_pos], xerr=time_err[tt], fmt='o', zorder=3)
            elif typ == 'preview':
                title = 'All Arrivals'

                if tt < divs[0]:
                    plt.annotate('F{:}'.format(tt + 1), (times[tt], y_pos), zorder=4)
                elif tt < divs[1] + divs[0]:
                    plt.annotate('B ', (times[tt], y_pos), zorder=4)
                elif tt < divs[2] + divs[1] + divs[0]:
                    plt.annotate('P{:}'.format(tt + 1 - divs[1] - divs[0]), (times[tt], y_pos), zorder=4)
                else:
                    plt.annotate('?', (times[tt], y_pos), zorder=4)
                
                plt.errorbar([times[tt]], [y_pos], xerr=time_err[tt], fmt='o', zorder=3)

        plt.xlabel('Time after {:} [s]'.format(self.bam.setup.fireball_datetime))
        plt.ylabel('Response')
        
        t_avg = np.nanmean(times)
        t_spr = np.nanstd(times)

        # I just REALLY want these to be nice graphs
        if t_spr > 0:
            plt.xlim(t_avg - SPREAD*t_spr, t_avg + SPREAD*t_spr)
        elif not np.isnan(t_avg):
            plt.xlim(t_avg - 5, t_avg + 5)
        elif not np.isnan(times[0]):
            plt.xlim(times[0] - 5, times[0] + 5)
        else:
            ref_pos = Position(self.bam.setup.lat_centre, self.bam.setup.lon_centre, 0)
            D = stn.stn_distance(ref_pos)
            plt.xlim(D/330 - 5, D/330 + 5)

        plt.title('{:} | Channel: {:}'.format(title, chn_selected))        

        pic_file = os.path.join(self.prefs.workdir, self.bam.setup.fireball_name, 'temp_Waveform.png')
        plt.savefig(pic_file)
        doc.add_picture(pic_file)
        os.remove(pic_file)
        plt.clf()

    def makeStaff(self, stn, ts, time_err, doc):

        # Generate pointes along trajecotry
        traj = self.bam.setup.trajectory
        points = traj.trajInterp2(div=100)

        D = stn.metadata.position
        ref_pos = Position(self.bam.setup.lat_centre, self.bam.setup.lon_centre, 0)
        D.pos_loc(ref_pos)

        times = []
        heights = []

        for pt in points:
            S = Position(pt[0], pt[1], pt[2])
            S.pos_loc(ref_pos)
            sounding, perts = self.bam.atmos.getSounding(lat=[S.lat, D.lat], lon=[S.lon, D.lon], heights=[S.elev, D.elev])

            f_time, _, _, _ = cyscan(S.xyz, D.xyz, sounding, wind=self.prefs.wind_en, n_theta=self.prefs.pso_theta, n_phi=self.prefs.pso_phi, \
                                    h_tol=self.prefs.pso_min_ang, v_tol=self.prefs.pso_min_dist)
            times.append(f_time + pt[3])
            heights.append(pt[2]/1000)

        plt.scatter(heights, times)

        X = np.linspace(17, 50)

        for tt in range(len(ts)):
            plt.fill_between(X, ts[tt]-time_err[tt][0], y2=ts[tt]+time_err[tt][1])
            plt.annotate('F{:}'.format(tt + 1), (18, ts[tt]))
            plt.axhline(y=ts[tt], color="black", linestyle="--")


        pic_file = os.path.join(self.prefs.workdir, self.bam.setup.fireball_name, 'staff.png')
        plt.savefig(pic_file)
        doc.add_picture(pic_file)
        os.remove(pic_file)
        plt.clf()

    def writeAtm(self, doc):


        lat = [self.bam.setup.lat_centre, self.bam.setup.lat_centre]
        lon = [self.bam.setup.lon_centre, self.bam.setup.lon_centre]
        elev = [50000, 0]

        sounding, perts = self.bam.atmos.getSounding(lat=lat, lon=lon, heights=elev)

        fig, axes = plt.subplots(nrows=1, ncols=3, sharey=True)
 

        def plotAx(ax, sounding, pert=False):


            h = sounding[:, 0]
            c = sounding[:, 1]
            mags = sounding[:, 2]
            dirs = np.degrees(sounding[:, 3])

            ax[0].plot(c, h)
            ax[0].set_xlabel('Sound Speed [m/s]')
            ax[0].set_ylabel('Height [m]')
            ax[2].plot(mags, h)
            ax[2].set_xlabel('Wind Speed [m/s]')
            ax[2].set_ylabel('Height [m]')
            ax[1].plot(dirs, h)
            ax[1].set_xlabel('Wind Direction [deg E from N]')
            ax[1].set_ylabel('Height [m]')



        plotAx(axes, sounding)
        for ptb in perts:
            plotAx(axes, ptb, pert=True)

        pic_file = os.path.join(self.prefs.workdir, self.bam.setup.fireball_name, 'atmos.png')
        plt.savefig(pic_file)
        doc.add_picture(pic_file)
        os.remove(pic_file)
        plt.clf()

    def writeSupSol(self, doc):
        
        doc.add_heading('Picks', 2)

        stat_picks_file_name = self.bam.setup.station_picks_file
        
        file_size = byteify(os.stat(stat_picks_file_name).st_size)
        
        doc.add_paragraph('Station Picks File: {:} ({:})'.format(stat_picks_file_name, file_size))

        with open(stat_picks_file_name, "r+") as f:

            lines = [line.rstrip().split(',') for line in f]

        rows = len(lines)
        cols = 8

        table = doc.add_table(rows=rows, cols=cols)
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Pick Group'
        header_cells[1].text = 'Network'
        header_cells[2].text = 'Code'
        header_cells[3].text = 'Latitude'
        header_cells[4].text = 'Longitude'
        header_cells[5].text = 'Elevation [m]'
        header_cells[6].text = 'Pick Time [s]'
        header_cells[7].text = 'Station Number'

        for ll, line in enumerate(lines):
            
            if ll == 0:
                continue
            
            row_cells = table.add_row().cells
            for i in range(cols):
                row_cells[i].text = str(line[i])
        
        results = supSearch(self.bam, self.prefs, manual=False, results_print=True, misfits=True)
        resultsPrint(results[0], results[1], results[2], results[3], self.prefs, doc=doc)

        pic_file = os.path.join(self.prefs.workdir, self.bam.setup.fireball_name, 'misfits_lat.png')        
        doc.add_picture(pic_file)
        os.remove(pic_file)

        pic_file = os.path.join(self.prefs.workdir, self.bam.setup.fireball_name, 'misfits_lon.png')        
        doc.add_picture(pic_file)
        os.remove(pic_file)

        pic_file = os.path.join(self.prefs.workdir, self.bam.setup.fireball_name, 'misfits_elev.png')        
        doc.add_picture(pic_file)
        os.remove(pic_file)

        pic_file = os.path.join(self.prefs.workdir, self.bam.setup.fireball_name, 'misfits_time.png')        
        doc.add_picture(pic_file)
        os.remove(pic_file)

        plt.clf()